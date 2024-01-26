import time
from tkinter.ttk import Combobox
import PyPDF2
import  os
import datetime
import pandas as pd
from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from docx import Document
from PIL import ImageDraw, ImageFont
from PIL import Image as PILImage
import textwrap
import pytz  # Importe a biblioteca pytz para lidar com fusos horários
import openpyxl
import json
from tkinter import *
import pandas as pd
import sys

# Redireciona a saída padrão para um arquivo
sys.stdout = open('output.log', 'w')
sys.stderr = open('error.log', 'w')

janela = Tk()

class Login():
    def __init__(self):
        self.janela = janela
        self.entry_user = None  # Variável que será acessada por outras classes
        self.formatação_aba()
        self.frame_principal()
        self.adicionar_titulo()
        self.botoes_tela_principal()
        janela.mainloop()

    
    def formatação_aba(self):
        self.janela.title("Carteira do Produtor")
        self.janela.configure(bg='aquamarine4')
        self.janela.geometry("600x300")
        self.janela.resizable(True, True)
        self.janela.minsize(width=300, height=250)

    def adicionar_titulo(self):
        # Adicionar título à tela principal
        titulo_label = Label(self.janela, text="Seu Login", font=('Arial', 12, 'bold'), bg='azure1', fg='black')
        titulo_label.place(relx=0.5, rely=0.05, anchor='n')

    def botoes_tela_principal(self):
        
        self.entry_user = Entry(self.frame0, bd=1, bg="white", highlightbackground="antiquewhite4", highlightthickness=1)
        self.entry_user.place(relx=0.19, rely=0.249, relwidth=0.6, relheight=0.1)
           # Vincular o evento Enter ao botão Confirmar
        self.entry_user.bind("<Return>", lambda event=None: self.exibir_tela_principal())

        self.bt_entrar = Button(self.frame0 , text="Entrar", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command=self.exibir_tela_principal)
        self.bt_entrar.place(relx=0.45, rely=0.505, relwidth=0.098, relheight=0.13)

        # Rótulo para exibir a mensagem de confirmação
        self.mensagem_label = Label(self.frame0, text="", font=('Arial', 10), bg='azure1', fg='green')
        self.mensagem_label.place(relx=0.5, rely=0.75, anchor='n')

    def frame_principal (self):
        self.frame0 = Frame(self.janela, bd=4, bg="azure1", highlightbackground="antiquewhite4", highlightthickness=3)
        self.frame0.place(relx=0.02, rely=0.02,relwidth=0.96, relheight=0.9)

    def exibir_tela_principal(self):
        usuario_digitado = self.entry_user.get()
        usuarios_aceitos = ["fjunior", "jose.miguel", "alexandramdasilva", "ssantana", "gbezerra", "lfernandes"]

        if usuario_digitado in usuarios_aceitos:
            self.frame0.destroy()  
            Application(self.janela, usuario_digitado)
        else:
            # Autenticação falhou
            self.mensagem_label.config(text="Usuário inválido. Tente novamente.", fg='red')
            self.entry_user.delete(0, END)


class Application:
    def __init__(self, janela, usuario_digitado):
        self.janela = janela
        self.usuario_digitado = usuario_digitado
        self.frame_principal()
        self.botoes_tela_principal()

    def botoes_tela_principal(self):
                # Adicionar título à tela principal
        titulo_label = Label(self.janela, text="Escolha uma opção:", font=('Arial', 12, 'bold'), bg='azure1', fg='black')
        titulo_label.place(relx=0.5, rely=0.05, anchor='n')
        
        self.bt_inserir_cpf = Button(self.frame1 , text="Inserir CPF", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command=self.exibir_tela_cpf)
        self.bt_inserir_cpf.place(relx=0.1, rely=0.2, relwidth=0.23, relheight=0.2)

        self.btemitir_cartao= Button(self.frame1 , text="Emitir Carteira", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command=self.exibir_tela_emitir_cartao)
        self.btemitir_cartao.place(relx=0.7, rely=0.2, relwidth=0.23, relheight=0.2)

        self.bt_fazer_memo = Button(self.frame1 , text="Fazer Memorando", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command=self.exibir_tela_memorando)
        self.bt_fazer_memo.place(relx=0.1, rely=0.6, relwidth=0.23, relheight=0.2)

        self.bttrocar_senha= Button(self.frame1 , text="Trocar Senha", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command=self.exibir_tela_trocar_senha)
        self.bttrocar_senha.place(relx=0.7, rely=0.6, relwidth=0.23, relheight=0.2)

    def frame_principal (self):
        self.frame1 = Frame(self.janela, bd=4, bg="azure1", highlightbackground="antiquewhite4", highlightthickness=3)
        self.frame1.place(relx=0.02, rely=0.02,relwidth=0.96, relheight=0.9)

    def exibir_tela_cpf(self):
        self.frame1.destroy()  
        formatacao_tela_cpf(self.janela, self.usuario_digitado)

    def exibir_tela_emitir_cartao(self):
        self.frame1.destroy()
        formatacao_tela_emitir_cartao(self.janela, self.usuario_digitado )
        
    def exibir_tela_memorando(self):
        self.frame1.destroy()
        formatacao_tela_emitir_memorando(self.janela, self.usuario_digitado)

    def exibir_tela_trocar_senha(self):
        self.frame1.destroy()
        formatacao_tela_trocar_senha(self.janela, self.usuario_digitado)

class formatacao_tela_emitir_cartao:

    def __init__(self, janela, usuario_digitado,):
        self.usuario_digitado = usuario_digitado
        self.botoes_tela_emitir_cartao()
        self.mensagem_var = StringVar()

    def carregar_senha(self):
        try:
            with open(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\senha.json") as file:
                data = json.load(file)
                return data.get("senha")
        except (FileNotFoundError, json.JSONDecodeError):
            return None

    def iniciar_automacao_web(self):
        options = webdriver.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--disable-notifications')
        options.add_argument('--window-size=1920,1080')
        options.add_argument('--log-level=3')  
        options.add_argument('--no-console') 
        options.add_argument('--disable-logging')

        self.driver = webdriver.Chrome(options=options)

        #Realizar login no site
        self.driver.get('http://sistemas.sefaz.am.gov.br/gcc/entrada.do')  # Substitua pelo URL do site
        usuario = self.driver.find_element(By.ID, 'username')  # Substitua pelo campo de usuário
        senha_da_pagina = self.driver.find_element(By.ID, 'password')  # Substitua pelo campo de senha
        botao_login = self.driver.find_element(By.XPATH, '//*[@id="fm1"]/fieldset/div[3]/div/div[4]/input[4]')  # Substitua pelo botão de login
    
        senha = self.carregar_senha()

        usuario.send_keys('')
        senha_da_pagina.send_keys(senha)
        botao_login.click()

    def fechar_driver(self):
        self.driver.quit()

    def frame_emitir_cartao(self):
        self.frame2 = Frame(janela, bd=4, bg="azure1", highlightbackground="antiquewhite4", highlightthickness=3)
        self.frame2.place(relx=0.02, rely=0.02, relwidth=0.96, relheight=0.9)

    def botoes_tela_emitir_cartao(self):
        self.janela = janela 
        self.frame_emitir_cartao()
        self.iniciar_automacao_web()

        titulo_label = Label(self.janela, text="Gerador de carteiras:", font=('Arial', 12, 'bold'), bg='azure1', fg='black')
        titulo_label.place(relx=0.5, rely=0.05, anchor='n')

        self.bt_fazer_impressao = Button(self.frame2, text="Gerar Carteiras", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command=self.fazer_carteiras)
        self.bt_fazer_impressao.place(relx=0.4, rely=0.19, relwidth=0.23, relheight=0.18)

        self.lista_cartoes = Listbox(self.frame2, bd=1, bg="white", highlightbackground="antiquewhite4", highlightthickness=1)
        self.lista_cartoes.place(relx=0.1, rely=0.4, relwidth=0.8, relheight=0.49)

        self.bt_voltar_principal = Button(self.frame2, text="Voltar", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command=self.voltar_tela_principal)
        self.bt_voltar_principal.place(relx=0.01, rely=0.01, relwidth=0.093, relheight=0.12)

    def fazer_carteiras(self):
           
        planilha = pd.read_excel(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\cpfs.xlsx")
        # Crie um DataFrame vazio para armazenar os dados
        # Crie um arquivo Excel para armazenar os dados
        os.path.isfile(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\banco_de_dados.xlsx")
        workbook = openpyxl.load_workbook(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\banco_de_dados.xlsx")
        sheet = workbook.active


        for index, row in planilha.iterrows():
            cpf = row['CPF']  # Pegar o CPF da linha atualself.

            # Configurar o fuso horário de Manaus (GMT-4)
            fuso_horario = pytz.timezone('America/Manaus')
            # Obtém a data atual no fuso horário de Manaus
            data_atual_manaus = datetime.datetime.now(fuso_horario)

            largura_maxima_x = 540  # Ajuste conforme necessário


            # Abrir o modelo PNG
            modelo = PILImage.open(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\img\frente.png")
            modelo_verso = PILImage.open(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\img\verso.png")


            # Carregar a fonte desejada (substitua pelo caminho da sua fonte)
            fonte = ImageFont.truetype(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\Roboto-Regular.ttf", 41)
            fonte_endereco = ImageFont.truetype(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\Roboto-Regular.ttf", 38)
            fonte_atv2 = ImageFont.truetype(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\Roboto-Regular.ttf", 39)

            # Criar um objeto de desenho
            desenho = ImageDraw.Draw(modelo)
            desenho_verso = ImageDraw.Draw(modelo_verso)

            while True:
                # Selecione o elemento usando o XPath
                elemento = self.driver.find_element(By.XPATH,'//*[@id="oCMenu___GCC2300"]')

                # Faça algo com o elemento, por exemplo, clique nele
                elemento.click()
                
                # Localize a lista suspensa usando o XPath fornecido (a <div> que representa o menu)
                dropdown_div = self.driver.find_element(By.XPATH, '//*[@id="oCMenu___GCC1008"]')

                # Clique na <div> para expandir o menu
                dropdown_div.click()
                # Suponha que você tenha o CPF que deseja colar em uma variável chamada "cpf"
                caixacpf = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cpfProdutorRuralFormatado"]')  # Substitua pelo XPath do campo de entrada
                caixacpf.click()  # Clique no campo de entrada para garantir que ele está ativo
                caixacpf.send_keys(cpf)  # Cole o conteúdo da área de transferência (ou 'cmd' em vez de 'ctrl' no Mac) 
                consultar = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_cadastroProdutorRuralAction!pesquisarProdutorRural"]')       
                consultar.click()
                time.sleep(1)
                situacao = self.driver.find_element(By.XPATH, '//*[@id="tbProdutorRural"]/thead/tr/th[2]')
                situacao.click()
                time.sleep(2)
                try:
                    abadeclaracao = WebDriverWait(self.driver, 10).until(
                        EC.presence_of_element_located((By.XPATH, '//*[@id="tbProdutorRural"]/tbody/tr[1]/td[8]/a[2]'))
                    )
                    abadeclaracao.click()
                    # Resto do seu código
                except Exception as e:
                    mensagem = (f"Erro ao clicar em 'abadeclaracao' para o CPF {cpf}: {str(e)}")

                #################################################################################
                nome_element = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cceaPessoaFisica_pfNome"]')
                # Use JavaScript para obter o valor do atributo 'value' do elemento
                nome = self.driver.execute_script("return arguments[0].value;", nome_element)

                rp_da_pagina = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_ieProdutorRuralFormatado"]')
                rp = self.driver.execute_script("return arguments[0]. value;", rp_da_pagina)

                cpf_element = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cpfProdutorRuralFormatado"]')
                cpf = self.driver.execute_script("return arguments[0]. value;", cpf_element)

                propiedade_element = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nmPropriedade"]')
                propriedade = self.driver.execute_script("return arguments[0]. value;", propiedade_element)

                endereco_element = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_txEnderecoPropriedade"]')
                endereco = self.driver.execute_script("return arguments[0]. value;", endereco_element)

                unloc_element = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_sgDistritoIdam"]')
                unloc_da_pagina = self.driver.execute_script("return arguments[0]. value;", unloc_element)
                # Mapeamento de valores numéricos para nomes de municípios
                ###########################
                if unloc_da_pagina == 'BAE' :
                    unloc_da_pagina = 'BAR'
                
                if unloc_da_pagina == "MTS-ATZ" or unloc_da_pagina =="MTS":
                    unloc_da_pagina = "ATZ-MTS"
                
                if unloc_da_pagina == "NRO-ITR" or unloc_da_pagina =="NRO":
                    unloc_da_pagina = "ITR-NRO"

                if unloc_da_pagina == "MTP-MNX" or unloc_da_pagina =="MTP":
                    unloc_da_pagina = "MNX-MTP"
                
                if unloc_da_pagina == "VE-LBR" or unloc_da_pagina =="VE":
                    unloc_da_pagina = "LBR-VE"

                if unloc_da_pagina == "VRC-MPU" or unloc_da_pagina =="VRC":
                    unloc_da_pagina = "MPU-VRC"
                    
                if unloc_da_pagina == "BNA-PRF" or unloc_da_pagina =="BNA":
                    unloc_da_pagina = "PRF-BNA"
                
                if unloc_da_pagina == "VLD-ITR" or unloc_da_pagina =="VLD":
                    unloc_da_pagina = "ITR-VLD"
                
                if unloc_da_pagina == "RLD-HIA" or unloc_da_pagina =="RLD":
                    unloc_da_pagina == "HIA-RLD"

                if unloc_da_pagina == "CAN-SUL":
                    unloc_da_pagina = "SUL-CAN"

                if unloc_da_pagina == "ZL-MAO" or unloc_da_pagina =="ZL":
                    unloc_da_pagina = "MAO-ZL"

                latitude_element = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_geoLatitude"]')
                latitude = self.driver.execute_script("return arguments[0]. value;", latitude_element)

                longitude_element =self. driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_geoLongitude"]')
                longitude = self.driver.execute_script("return arguments[0]. value;", longitude_element)

                atv1_element = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nmCnaePrincipal"]')
                atv1= self.driver.execute_script("return arguments[0]. value;", atv1_element)


                atv2_element = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nmCnaeSecundario"]')
                atv2 = self.driver.execute_script("return arguments[0]. value;", atv2_element)

                inicioatv_element = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_anoInicioAtividade"]')
                inicioatv = self.driver.execute_script("return arguments[0]. value;", inicioatv_element)

                numcontrole_element = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_nrDeclaracaoUnidLocal"]')
                numcontrole = self.driver.execute_script("return arguments[0]. value;",numcontrole_element)

                cnae1_element = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cnaePrincipalFormatado"]')
                cnae1 = self.driver.execute_script("return arguments[0]. value;",cnae1_element)

                cnae2_element = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cnaeSecundarioFormatado"]')
                cnae2 = self.driver.execute_script("return arguments[0]. value;",cnae2_element)

                descricao = self.usuario_digitado
                
                if not atv2:
                    cnae2 = ""
            
                # Após coletar as informações
                data_atual_manaus_str = data_atual_manaus.strftime("%d/%m/%Y")
                nova_linha = [nome, cpf, unloc_da_pagina, data_atual_manaus_str, cnae1, cnae2, descricao]
                sheet.append(nova_linha)
                
                unloc = "PR-" + unloc_da_pagina + "/" + numcontrole
                unloc = str(unloc)
                validade = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_dtValidadeDeclaracaoFormatado"]')
                validade_da_pagina = self.driver.execute_script("return arguments[0]. value;",validade)
                mensagem = (nome, "|", unloc, "|")
            # Verifique se a pasta correspondente à unloc já existe
                pasta_unloc = os.path.join(r"Y:\CARTEIRAS DIGITAIS", unloc_da_pagina)
                
                if not os.path.exists(pasta_unloc):
                    # Se a pasta não existe, crie-a
                    os.mkdir(pasta_unloc)

                # Caminho completo para salvar o PDF na pasta correspondente crie uma pasta com o dia atual (formato: YYYY-MM-DD)
                data_atual = datetime.datetime.now()
                pasta_data_atual = os.path.join(pasta_unloc, data_atual.strftime("%d.%m.%Y"))

                if not os.path.exists(pasta_data_atual):
                    # Se a pasta da data atual não existe, crie-a
                    os.mkdir(pasta_data_atual)
                        
                # Caminho completo para salvar o PDF na pasta da data atual
                pdf_path = os.path.join(pasta_data_atual, nome + '.pdf')

                # Defina a pasta da data atual como a pasta para salvar o PDF
                output_pdf_filename = pdf_path

                # Definir as coordenadas para colar as informações (x, y)
                coordenadas_frente = {
                    "RP": (217,393),
                    "NOME": (95,518,1309,604),
                    "CPF": (864,392),
                    "NOME PROPRIEDADE": (100, 660),
                    "UNLOC": (212,824),
                    "INICIO ATV": (751,824),
                    "VALIDADE": (1063,825),
                }

                # Inserir informações no modelo nas coordenadas especificadas
                desenho.text(coordenadas_frente["RP"], rp, fill=(0,0,0), font=fonte)
                desenho.text(coordenadas_frente["NOME"], nome, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["CPF"], cpf, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["NOME PROPRIEDADE"], propriedade, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["UNLOC"], unloc, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["INICIO ATV"], inicioatv, fill=(0, 0, 0), font=fonte)
                desenho.text(coordenadas_frente["VALIDADE"], validade_da_pagina, fill=(0, 0, 0), font=fonte)

                # Quebra o texto em várias linhas
                linhas_endereco = textwrap.wrap(endereco, largura_maxima_x)
                def limitar_texto(texto, comprimento_maximo):
                    if len(texto) > comprimento_maximo:
                            return texto[:comprimento_maximo - 3]
                    else:
                        return texto
                endereco = limitar_texto(endereco, 50)

                linhas_atv2 = textwrap.wrap(atv2, largura_maxima_x)
                def limitar_texto(texto, comprimento_maximo):
                    if len(texto) > comprimento_maximo:
                            return texto[:comprimento_maximo - 3]
                    else:
                        return texto
                atv2 = limitar_texto(atv2, 60)
                
                # Inserir informações no modelo nas coordenadas especificadas
                coordenadas_verso = {
                    "END": (89,285),
                    "ATV1": (89,473),
                    "ATV2": (89,631),
                    "LOC": (382,804),
                }

                def desenhar_texto_quebrado(coordenadas, texto, desenho, fonte, largura_maxima_x):
                    coordenada = (coordenadas[0], coordenadas[1])
                    linhas = textwrap.wrap(texto, width=largura_maxima_x // 9)
                    for linha in linhas:
                        linhas_quebradas = textwrap.wrap(linha, width=largura_maxima_x // 9)

                        for linha_quebrada in linhas_quebradas:
                            if coordenada[1] < modelo_verso.size[1]:
                                desenho.text(coordenada, linha_quebrada, fill=(0, 0, 0), font=fonte)
                                coordenada = (coordenada[0], coordenada[1] + 40)  # Aumente o valor (ex: +50) conforme necessário
                            else:
                                break

                desenho_verso.text(coordenadas_verso["END"], endereco, fill=(0, 0, 0), font=fonte_endereco)

                coordenada_endereco = (coordenadas_verso["END"][0], coordenadas_verso["END"][1])

                # Loop para desenhar cada linha do endereço
                for linha in linhas_endereco:
                    linhas_quebradas = textwrap.wrap(linha, width=largura_maxima_x // 9)

                    for linha_quebrada in linhas_quebradas:
                        if coordenada_endereco[1] < modelo_verso.size[1]:
                            desenho_verso.text(coordenada_endereco, linha_quebrada, fill=(0, 0, 0), font=fonte_endereco)
                            coordenada_endereco = (coordenada_endereco[0], coordenada_endereco[1] + 40)  # Aumente o valor (ex: +50) conforme necessário
                        else:
                            break


                desenhar_texto_quebrado(coordenadas_verso["END"], endereco, desenho_verso, fonte_endereco, largura_maxima_x)
                desenhar_texto_quebrado(coordenadas_verso["ATV1"], cnae1 + " - " + atv1, desenho_verso, fonte, largura_maxima_x)
                if atv2:
                    desenhar_texto_quebrado(coordenadas_verso["ATV2"], cnae2 + " - " + atv2, desenho_verso, fonte, largura_maxima_x)
                desenhar_texto_quebrado(coordenadas_verso["LOC"], latitude + "  " + longitude, desenho_verso, fonte, largura_maxima_x)

                # Salvar o novo PNG
                modelo.save(cpf + ".pdf")
                modelo_verso.save("verso.pdf")
                # Nome dos arquivos PDF que você gerou
                pdf1_filename = cpf + ".pdf"
                pdf2_filename = "verso.pdf"

                # Abra os arquivos PDF
                pdf1 = PyPDF2.PdfReader(pdf1_filename)
                pdf2 = PyPDF2.PdfReader(pdf2_filename)

                # Crie um objeto PDFFileMerger para mesclar os PDFs
                pdf_merger = PyPDF2.PdfMerger()

                # Adicione os arquivos PDF à mesclagem
                pdf_merger.append(pdf1)
                pdf_merger.append(pdf2)

                output_pdf_filename = os.path.join(r"Y:\CARTEIRAS DIGITAIS",unloc_da_pagina, pasta_data_atual , nome  + '.pdf')
                pdf_merger.write(output_pdf_filename)

                # Feche o arquivo PDF de saída
                pdf_merger.close()

                # Excluir os arquivos originais (frente e verso)
                os.remove(cpf + ".pdf")
                os.remove("verso.pdf")

                # Após a geração do PDF com sucesso, você pode remover a linha correspondente ao CPF do DataFrame
                planilha = planilha[planilha['CPF'] != cpf]
                        # Adicione a mensagem à lista
                        # Adicione a mensagem à lista_cartoes imediatamente após a conclusão do processo para cada cartão
                self.lista_cartoes.insert(END, mensagem)

                break
            
        workbook.save(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\banco_de_dados.xlsx")
        # Salvar a planilha atualizada após a exclusão do CPF
        planilha.to_excel(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\cpfs.xlsx", index=False)

    def voltar_tela_principal(self):
        self.fechar_driver()  # Feche o driver Selenium, se necessário
        self.frame2.destroy()  # Destrua a tela atual
        Application(janela, self.usuario_digitado)  # Crie novamente a tela principal


class formatacao_tela_cpf:

    def __init__(self, janela, usuario_digitado):
        self.usuario_digitado = usuario_digitado
        self.botoes_tela_cpf()
        self.mensagem_var = StringVar()

    def formatar_cpf(self, cpf):
        cpf_formatado = '{}.{}.{}-{}'.format(cpf[:3], cpf[3:6], cpf[6:9], cpf[9:])
        return cpf_formatado
    
    def carregar_senha(self):
        try:
            with open(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\senha.json") as file:
                data = json.load(file)
                return data.get("senha")
        except (FileNotFoundError, json.JSONDecodeError):
            return None

    def iniciar_automacao_web(self):
        options = webdriver.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--disable-notifications')
        options.add_argument('--window-size=1920,1080')
        options.add_argument('--log-level=3')  
        options.add_argument('--no-console') 
        options.add_argument('--disable-logging')

        self.driver = webdriver.Chrome(options=options)

        #Realizar login no site
        self.driver.get('http://sistemas.sefaz.am.gov.br/gcc/entrada.do')  # Substitua pelo URL do site
        usuario = self.driver.find_element(By.ID, 'username')  # Substitua pelo campo de usuário
        senha_da_pagina = self.driver.find_element(By.ID, 'password')  # Substitua pelo campo de senha
        botao_login = self.driver.find_element(By.XPATH, '//*[@id="fm1"]/fieldset/div[3]/div/div[4]/input[4]')  # Substitua pelo botão de login
    
        senha = self.carregar_senha()

        usuario.send_keys('')
        senha_da_pagina.send_keys(senha)
        botao_login.click()

        
    def fechar_driver(self):
        self.driver.quit()

    def frame_cpf(self):
        self.frame2 = Frame(janela, bd=4, bg="azure1", highlightbackground="antiquewhite4", highlightthickness=3)
        self.frame2.place(relx=0.02, rely=0.02, relwidth=0.96, relheight=0.9)

    def botoes_tela_cpf(self):
        self.janela = janela
        self.frame_cpf()
        self.iniciar_automacao_web()

        titulo_label = Label(self.janela, text="Inserir CPF:", font=('Arial', 12, 'bold'), bg='azure1', fg='black')
        titulo_label.place(relx=0.5, rely=0.05, anchor='n')

        self.entry_cpf = Entry(self.frame2, bd=1, bg="white", highlightbackground="antiquewhite4", highlightthickness=1)
        self.entry_cpf.place(relx=0.1, rely=0.249, relwidth=0.6, relheight=0.1)
           # Vincular o evento Enter ao botão Confirmar
        self.entry_cpf.bind("<Return>", lambda event=None: self.consulta_cpfs())

        self.bt_confirmar = Button(self.frame2, text="Confirmar", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command=self.consulta_cpfs)
        self.bt_confirmar.place(relx=0.72, rely=0.2, relwidth=0.23, relheight=0.18)

        titulo2_label = Label(self.janela, text="Status:", font=('Arial', 10, 'italic'), bg='azure1', fg='black')
        titulo2_label.place(relx=0.5, rely=0.45, anchor='n')

        self.lista_cpfs = Listbox(self.frame2, bd=1, bg="white", highlightbackground="antiquewhite4", highlightthickness=1)
        self.lista_cpfs.place(relx=0.1, rely=0.6, relwidth=0.8, relheight=0.3)

        self.bt_voltar_principal = Button(self.frame2, text="Voltar", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command=self.voltar_tela_principal)
        self.bt_voltar_principal.place(relx=0.01, rely=0.01, relwidth=0.093, relheight=0.12)

    def consulta_cpfs(self):
        planilha = pd.read_excel(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\cpfs.xlsx")
        cpf_digitado = self.entry_cpf.get()
        if cpf_digitado:
            cpf_formatado = self.formatar_cpf(cpf_digitado)
            # Verifica se o CPF já existe na planilha
            if cpf_formatado in planilha['CPF'].values:
                mensagem = f"CPF {cpf_formatado} já está na fila de impressão."
            else:
                # Verificar se o CPF já existe na planilha
                elemento = self.driver.find_element(By.XPATH,'//*[@id="oCMenu___GCC2300"]')
                elemento.click()

                dropdown_div = self.driver.find_element(By.XPATH, '//*[@id="oCMenu___GCC1008"]')
                dropdown_div.click()

                caixacpf = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_produtorRuralTOA_produtorRural_cpfProdutorRuralFormatado"]')
                caixacpf.click()
                caixacpf.send_keys(cpf_formatado)

                consultar = self.driver.find_element(By.XPATH, '//*[@id="formProdutorRural_cadastroProdutorRuralAction!pesquisarProdutorRural"]')
                consultar.click()

                try:
                    comcpf = self.driver.find_element(By.XPATH, '//*[@id="tbProdutorRural"]/tbody/tr/td[8]/a[2]/img')
                    planilha = pd.read_excel(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\cpfs.xlsx")
                    novo_dataframe = pd.DataFrame({'CPF': [cpf_formatado]})
                    planilha = pd.concat([planilha, novo_dataframe], ignore_index=True)
                    planilha.to_excel(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\cpfs.xlsx", index=False)
                    mensagem = f"CPF {cpf_formatado} inserido com sucesso!"

                except NoSuchElementException:
                    mensagem = f"CPF {cpf_formatado} não encontrado"

              
            self.lista_cpfs.insert(END, mensagem)
            self.entry_cpf.delete(0, END)  # Limpa a entrada do CPF

    def voltar_tela_principal(self):
        self.fechar_driver()  # Feche o driver Selenium, se necessário
        self.frame2.destroy()  # Destrua a tela atual
        Application(janela, self.usuario_digitado)  # Crie novamente a tela principal
        
class formatacao_tela_emitir_memorando:

    def __init__(self, janela, usuario_digitado):
        self.usuario_digitado = usuario_digitado
        self.botoes_tela_memorando()
        self.formatação_aba_memo()
        self.mensagem_var = StringVar()
    
    def formatação_aba_memo(self):
        self.janela.geometry("800x500")

    def frame_memorando(self):
        self.frame3 = Frame(janela, bd=4, bg="azure1", highlightbackground="antiquewhite4", highlightthickness=3)
        self.frame3.place(relx=0.02, rely=0.02, relwidth=0.96, relheight=0.9)

    def botoes_tela_memorando(self):
        self.janela = janela
        self.frame_memorando()

        self.titulo_label = Label(self.janela, text="Insira as Informações:", font=('Arial', 12, 'bold'), bg='azure1', fg='black')
        self.titulo_label.place(relx=0.5, rely=0.05, anchor='n')

        self.memorando = Entry(self.frame3, bd=1, bg="white", highlightbackground="antiquewhite4", highlightthickness=1)
        self.memorando.place(relx=0.01, rely=0.249, relwidth=0.4, relheight=0.055)

        self.municipio = Entry(self.frame3, bd=1, bg="white", highlightbackground="antiquewhite4", highlightthickness=1)
        self.municipio.place(relx=0.01, rely=0.4, relwidth=0.4, relheight=0.055)

        self.data_emissao = Entry(self.frame3, bd=1, bg="white", highlightbackground="antiquewhite4", highlightthickness=1)
        self.data_emissao.place(relx=0.6, rely=0.249, relwidth=0.4, relheight=0.055)

        self.memos_utilizados = Entry(self.frame3, bd=1, bg="white", highlightbackground="antiquewhite4", highlightthickness=1)
        self.memos_utilizados.place(relx=0.6, rely=0.4, relwidth=0.4, relheight=0.055)


        nmemo_label = Label(self.janela, text="N° do Memorando:", font=('Arial', 9,), bg='azure1', fg='black')
        nmemo_label.place(relx=0.2, rely=0.180, anchor='n')

          # Utilizando o Combobox para o município
        self.municipio_var = StringVar()
        municipios = ['ALV', 'AMT', 'ANA', 'ANO', 'APU', 'ATN', 'ATZ','ATZ-MTS', 'BAE', 'BAZ', 'BAR', 'BBA', 'BVR', 'BER', 'BJC', 'BOA', 'CAF', 'CAN', 'CAP', 'CAR', 'CAZ', 'CIZ', 'COD', 'CRZ', 'ENV', 'ERN', 'FBA', 'GAJ', 'HIA','HIA-RLD', 'IRB', 'ITA', 'ITG', 'ITR', 'ITR-NRO', 'ITR-VLD', 'IPX', 'JPR', 'JUR', 'JUT', 'LBR', 'LBR-VE', 'MAO','MAO-ZL', 'MBZ', 'MNX', 'MNX-MTP', 'MRA',  'MPU', 'MPU-VRC', 'NAP', 'NAR', 'NMD', 'NON', 'PAR', 'PUI', 'PRF', 'PRF-BNA',  'RPE', 'SAI', 'SJL', 'SPO', 'SSU', 'SUL-CAN', 'SLV', 'TBT', 'TPA', 'TFF', 'TNT', 'UAN', 'UCB', 'URC']
  # Substitua com suas siglas desejadas
        self.municipio_combobox = Combobox(self.frame3, textvariable=self.municipio_var, values=municipios, state="readonly")
        self.municipio_combobox.place(relx=0.01, rely=0.4, relwidth=0.4, relheight=0.055)


        data_emissao = Label(self.janela, text="Data de emissão da carteira:", font=('Arial', 9,), bg='azure1', fg='black')
        data_emissao.place(relx=0.8, rely=0.180, anchor='n')

        memos_utilizados = Label(self.janela, text="Memorando Utilizado:", font=('Arial', 9,), bg='azure1', fg='black')
        memos_utilizados.place(relx=0.8, rely=0.33, anchor='n')

        self.btenviar= Button(self.frame3 , text="Enviar", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command= self.automacao_memorando)
        self.btenviar.place(relx=0.45, rely=0.505, relwidth=0.073, relheight=0.05)

        self.bt_voltar_principal = Button(self.frame3, text="Voltar", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command=self.voltar_tela_principal)
        self.bt_voltar_principal.place(relx=0.01, rely=0.01, relwidth=0.073, relheight=0.085)

        self.confirmacao = Listbox(self.frame3, bd=1, bg="white", highlightbackground="antiquewhite4", highlightthickness=1)
        self.confirmacao.place(relx=0.1, rely=0.6, relwidth=0.8, relheight=0.3)
   
    def automacao_memorando(self):
        from datetime import datetime
        mapeamento_municipios = {
                'ALV': 'Alvarães',
                'AMT': 'Amaturá',
                'ANA':'Anamã',
                'ANO':'Anori',          
                'APU':'Apuí',
                'ATN':'Atalaia do Norte',
                'ATZ':'Autazes',
                'BAZ':'Barcelos',
                'BAR':'Barreirinha',
                'BJC':'Benjamin Constant',
                'BER':'Beruri',
                'BVR':'Boa Vista do Ramos',
                'BOA':'Boca do Acre',
                'BBA':'Borba',
                'CAP':'Caapiranga',
                'CAN':'Canutama',
                'CAF':'Carauari',
                'CAR':'Careiro',
                'CAZ':'Careiro da Várzea',
                'CIZ':'Coari',
                'COD':'Codajás',
                'ERN':'Eirunepé',
                'ENV':'Envira',
                'FBA':'Fonte Boa',
                'GAJ':'Guajará',
                'HIA':'Humaitá',
                'IPX':'Ipixuna',
                'IRB':'Iranduba',
                'ITA':'Itamarati',
                'ITR':'Itacoatiara',
                'ITG':'Itapiranga',
                'JPR':'Japurá',
                'JUR':'Juruá',
                'JUT':'Jutaí',
                'LBR':'Lábrea',
                'MPU':'Manacapuru',
                'MQR':'Manaquiri',
                'MAO':'Manaus',
                'MAO-ZL':'Manaus ZL',
                'MNX':'Manicoré',
                'ATZ-MTS':'Monte Sinai',
                'MRA':'Maraã',
                'MBZ':'Maués',
                'NMD':'Nhamundá',
                'ITR-NRO':'Novo Remanso',
                'MNX-MTP':'Santo Antonio do Matupi',
                'LBR-VE':'Vila Extrema',
                'MPU-VRC':'Vila Rica de Caviana',
                'PRF-BNA':'Balbina',
                'ITR-VLD':'Vila de Lindoia',
                'RLD-HIA':'Vila da Realidade',
                'NON':'Nova Olinda do Norte',
                'NAR':'Novo Airão',
                'NAP':'Novo Aripuanã',
                'PAR':'Parintins',
                'PUI':'Pauini',
                'PRF':'Presidente Figueiredo',
                'RPE':'Rio Preto da Eva',
                'SIR':'Santa Isabel do Rio Negro',
                'SAI':'Santo Antônio do Içá',
                'SJL':'São Gabriel da Cachoeira',
                'SPO':'São Paulo de Olivença',
                'SSU':'São Sebastião do Uatumã',
                'SUL-CAN':'Sul de Canutama',
                'SLV':'Silves',
                'TBT':'Tabatinga',
                'TPA':'Tapauá',
                'TFF':'Tefé',
                'TNT':'Tonantins',
                'UAN':'Uarini',
                'URC':'Urucará',
                'UCB':'Urucurituba',
                'MAO-ZL': 'Manaus-zl'
            }

        # Abre o arquivo Excel
        workbook = openpyxl.load_workbook(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\banco_de_dados.xlsx")

        # Carregar o modelo do Word
        doc = Document(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\MemorandoSaida\modelo_memosaida.docx")

        # Carregar os dados do Excel em um DataFrame
        df = pd.read_excel(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\banco_de_dados.xlsx")

        
        # Obter entrada do usuário para o município e a data desejados
        memorando = self.memorando.get()
        unloc = self.municipio_combobox.get()
        data = self.data_emissao.get()
        memo = self.memos_utilizados.get()
    # Filtrar os dados do DataFrame com base na entrada do usuário
        selected_data = df[(df['municipio'] == unloc) & (df['data'] == data)]

                    # Obtendo a UNLOC correspondente ao município
        municipio = mapeamento_municipios.get(unloc, '')


        # Verificar se existem dados para o município e a data especificados
        if selected_data.empty:
            mensagem = (f"Não foram encontrados dados para o município '{unloc}' na data '{data}'.")
        else:
            # Ordenar a coluna 'nomes' em ordem alfabética
            selected_data = selected_data.sort_values(by='nomes')
            # Preencher o modelo do Word com as informações coletadas e adicionar numeração
            for paragraph in doc.paragraphs:
                text = paragraph.text

                if "(num)" in text:
                    # Adicionar a data formatada somente se a tag "(data)" estiver presente
                    text = text.replace("(num)", memorando)
                
                # Verificar se a tag "(data)" está presente no texto
                if "(data)" in text:
                    # Adicionar a data formatada somente se a tag "(data)" estiver presente
                    text = text.replace("(data)", data)
                
                # Verificar se a tag "(muni)" está presente no texto
                if "(muni)" in text:
                    # Substituir o nome do município e aplicar formatação em negrito
                    text = text.replace("(muni)", municipio)
                    run = paragraph.add_run(municipio)
                    run.bold = True  # Colocar o texto em negrito
                    paragraph.clear()

                if "(memos)" in text:
                    # Adicionar a data formatada somente se a tag "(data)" estiver presente
                    text = text.replace("(memos)", memo)
                
                nomes = selected_data['nomes'].tolist()
                # Adicionar numeração antes de cada nome
                nomes_numerados = [f"{i + 1}. {nome}" for i, nome in enumerate(nomes)]
                text = text.replace("(nomes)", '\n'.join(nomes_numerados))

                
                # Obter o total de nomes (último valor da enumeração) e substituir na variável 'qtda'
                qtda = len(nomes_numerados)
                text = text.replace("(qtda)", str(qtda))
                
                paragraph.clear()  # Limpar o parágrafo original
                paragraph.add_run(text)  # Adicionar o texto modificado
                # Criação da mensagem com quebra de linha entre os nomes
                mensagem = 'NOMES INSERIDOS: \n{}'.format('\n'.join(nomes))
            # Salvar o documento preenchido na pasta desejada
            output_path = r'Y:\MEMORANDOS CPCPR\MEMORANDO 2024'
            doc.save(output_path + '\\' + memorando + '-' + municipio + '-' + 'DIG' + '.docx')


        # Verificar se a UNLOC correspondente foi encontrada
        if municipio:
            # Verifique se a pasta correspondente à UNLOC já existe
            pasta_unloc = os.path.join(r"Y:\CARTEIRAS DIGITAIS", unloc)

            # Crie uma pasta com o dia atual (formato: YYYY-MM-DD)
            data_formatada = datetime.strptime(data, "%d/%m/%Y").strftime("%d.%m.%Y")

            # Caminho completo para salvar o documento preenchido na pasta da data atual
            output_path = os.path.join(pasta_unloc, data_formatada, memorando + '-' + municipio + '-' + 'DIG' + '.docx')

            # Salvar o documento preenchido na pasa correspondente
            doc.save(output_path)
        else:
            mensagem = (f"A UNLOC correspondente ao município '{municipio}' não foi encontrada.")

        # Fechar o arquivo Excel
        workbook.close()
        self.memorando.delete(0, END) 
        self.municipio.delete(0, END) 
        self.data_emissao.delete(0, END) 
        self.memos_utilizados.delete(0, END)
        self.confirmacao.insert(END, mensagem) 

    def formatação_aba(self):
        self.janela.title("Carteira do Produtor")
        self.janela.configure(bg='aquamarine4')
        self.janela.geometry("600x300")
        self.janela.resizable(True, True)
        self.janela.minsize(width=300, height=250)

    def voltar_tela_principal(self):
        self.frame3.destroy()  # Destrua a tela atual
        self.formatação_aba()
        Application(janela, self.usuario_digitado)  # Crie novamente a tela principal

class formatacao_tela_trocar_senha:
    def __init__(self, janela, usuario_digitado):
        self.usuario_digitado = usuario_digitado
        self.janela = janela
        self.botoes_tela_trocar_senha()

    def frame_trocar_senha(self):
        self.frame4 = Frame(self.janela, bd=4, bg="azure1", highlightbackground="antiquewhite4", highlightthickness=3)
        self.frame4.place(relx=0.02, rely=0.02, relwidth=0.96, relheight=0.9)

    def carregar_senha(self):
        try:
            with open(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\senha.json") as file:
                data = json.load(file)
                return data.get("senha")
        except (FileNotFoundError, json.JSONDecodeError):
            return None

    def botoes_tela_trocar_senha(self):
        self.frame_trocar_senha()

        titulo_label = Label(self.janela, text="Insira a Nova Senha", font=('Arial', 12, 'bold'), bg='azure1', fg='black')
        titulo_label.place(relx=0.5, rely=0.05, anchor='n')

        self.nova_senha_var = StringVar()
        self.nova_senha = Entry(self.frame4, bd=1, bg="white", highlightbackground="antiquewhite4", highlightthickness=1, textvariable=self.nova_senha_var)
        self.nova_senha.place(relx=0.1, rely=0.249, relwidth=0.6, relheight=0.1)

        # Vincular o evento Enter ao botão Confirmar
        self.nova_senha.bind("<Return>", lambda event=None: self.salvar_nova_senha())

        self.bt_confirmar = Button(self.frame4, text="Confirmar", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command=self.salvar_nova_senha)
        self.bt_confirmar.place(relx=0.72, rely=0.2, relwidth=0.23, relheight=0.18)

        self.bt_voltar_principal = Button(self.frame4, text="Voltar", bd=4, bg="white", highlightbackground="antiquewhite4", highlightthickness=3, command=self.voltar_tela_principal)
        self.bt_voltar_principal.place(relx=0.01, rely=0.01, relwidth=0.093, relheight=0.12)

    def salvar_nova_senha(self):
        nova_senha_valor = self.nova_senha_var.get()
        data = {"senha": nova_senha_valor}
        
        with open(r"Y:\ARQUIVO DIGITAL CPR\fabio jr\CARTAO-DIGITAL\senha.json", "w") as file:
            json.dump(data, file)

        self.nova_senha.delete(0, END)  # Limpa a entrada do CPF

    def voltar_tela_principal(self):
        self.frame4.destroy()  # Destrua a tela atual
        Application(janela, self.usuario_digitado)  # Crie novamente a tela principal

if __name__ == "__main__":
    Login()
